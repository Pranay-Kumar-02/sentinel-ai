import os
import re
import io
import base64
import tempfile
from pathlib import Path
import pytesseract
from PIL import Image
import cv2
import numpy as np

# Set Tesseract path for Windows
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# ── URL Extraction Helper ──────────────────────────────────────────────────────
def extract_urls_from_text(text: str) -> list:
    pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
    urls = re.findall(pattern, text)
    pattern2 = r'(?:www\.)[^\s<>"{}|\\^`\[\]]+'
    urls2 = re.findall(pattern2, text)
    return list(set(urls + urls2))

# ── SCREENSHOT / IMAGE ANALYSIS ────────────────────────────────────────────────
async def analyze_image(file_bytes: bytes, filename: str) -> dict:
    """
    Extract text from screenshot/image using OCR then analyze for threats.
    """
    try:
        # Convert bytes to PIL Image
        image = Image.open(io.BytesIO(file_bytes))

        # Convert to RGB if needed
        if image.mode not in ("RGB", "L"):
            image = image.convert("RGB")

        # Enhance image for better OCR
        img_array = np.array(image)
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)

        # Apply thresholding to improve OCR accuracy
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        enhanced = Image.fromarray(thresh)

        # Run OCR
        extracted_text = pytesseract.image_to_string(enhanced, config='--psm 6')
        extracted_text = extracted_text.strip()

        if not extracted_text:
            # Try without enhancement
            extracted_text = pytesseract.image_to_string(image)
            extracted_text = extracted_text.strip()

        # Extract URLs from OCR text
        urls = extract_urls_from_text(extracted_text)

        # Image metadata
        width, height = image.size

        return {
            "file_type":       "image",
            "filename":        filename,
            "dimensions":      f"{width}x{height}px",
            "extracted_text":  extracted_text,
            "extracted_urls":  urls,
            "char_count":      len(extracted_text),
            "ocr_success":     len(extracted_text) > 10,
            "ready_for_analysis": len(extracted_text) > 5,
        }

    except Exception as e:
        return {
            "file_type":   "image",
            "filename":    filename,
            "error":       str(e),
            "extracted_text": "",
            "extracted_urls": [],
            "ready_for_analysis": False,
        }


# ── QR CODE ANALYSIS ───────────────────────────────────────────────────────────
async def analyze_qr(file_bytes: bytes, filename: str) -> dict:
    """
    Decode QR code from image and extract the URL/data inside.
    """
    try:
        from pyzbar.pyzbar import decode as qr_decode

        # Convert to OpenCV image
        img_array = np.frombuffer(file_bytes, np.uint8)
        img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

        if img is None:
            return {"error": "Could not read image", "qr_data": None, "file_type": "qr"}

        # Try to decode QR
        decoded = qr_decode(img)

        if not decoded:
            # Try grayscale
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            decoded = qr_decode(gray)

        if not decoded:
            # Try with enhanced contrast
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            enhanced = clahe.apply(cv2.cvtColor(img, cv2.COLOR_BGR2GRAY))
            decoded = qr_decode(enhanced)

        if decoded:
            qr_data = decoded[0].data.decode("utf-8")
            qr_type = decoded[0].type

            return {
                "file_type":    "qr",
                "filename":     filename,
                "qr_data":      qr_data,
                "qr_type":      qr_type,
                "is_url":       qr_data.startswith("http") or qr_data.startswith("www"),
                "extracted_urls": [qr_data] if qr_data.startswith("http") else [],
                "extracted_text": qr_data,
                "ready_for_analysis": True,
            }
        else:
            return {
                "file_type":    "qr",
                "filename":     filename,
                "error":        "No QR code detected in image",
                "qr_data":      None,
                "extracted_urls": [],
                "ready_for_analysis": False,
            }

    except Exception as e:
        return {
            "file_type":   "qr",
            "filename":    filename,
            "error":       str(e),
            "qr_data":     None,
            "extracted_urls": [],
            "ready_for_analysis": False,
        }


# ── PDF ANALYSIS ───────────────────────────────────────────────────────────────
async def analyze_pdf(file_bytes: bytes, filename: str) -> dict:
    """
    Extract text, links, and metadata from PDF files.
    """
    try:
        import fitz  # PyMuPDF

        pdf = fitz.open(stream=file_bytes, filetype="pdf")

        all_text = ""
        all_urls = []
        page_count = len(pdf)
        metadata = pdf.metadata or {}

        for page_num in range(page_count):
            page = pdf[page_num]

            # Extract text
            text = page.get_text()
            all_text += text + "\n"

            # Extract embedded links
            links = page.get_links()
            for link in links:
                if link.get("uri"):
                    all_urls.append(link["uri"])

        # Also extract URLs from text
        text_urls = extract_urls_from_text(all_text)
        all_urls = list(set(all_urls + text_urls))

        pdf.close()

        return {
            "file_type":       "pdf",
            "filename":        filename,
            "page_count":      page_count,
            "extracted_text":  all_text.strip()[:5000],  # Limit to 5000 chars
            "extracted_urls":  all_urls,
            "url_count":       len(all_urls),
            "char_count":      len(all_text),
            "metadata": {
                "title":    metadata.get("title", ""),
                "author":   metadata.get("author", ""),
                "creator":  metadata.get("creator", ""),
                "producer": metadata.get("producer", ""),
            },
            "ready_for_analysis": len(all_text.strip()) > 10 or len(all_urls) > 0,
        }

    except Exception as e:
        return {
            "file_type":   "pdf",
            "filename":    filename,
            "error":       str(e),
            "extracted_text": "",
            "extracted_urls": [],
            "ready_for_analysis": False,
        }


# ── DOCX ANALYSIS ─────────────────────────────────────────────────────────────
async def analyze_docx(file_bytes: bytes, filename: str) -> dict:
    """
    Extract text and hyperlinks from Word documents.
    """
    try:
        from docx import Document
        import io

        doc = Document(io.BytesIO(file_bytes))

        all_text = ""
        all_urls = []

        # Extract paragraph text
        for para in doc.paragraphs:
            all_text += para.text + "\n"

        # Extract hyperlinks from relationships
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                all_urls.append(rel.target_ref)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "

        # Extract URLs from text
        text_urls = extract_urls_from_text(all_text)
        all_urls = list(set(all_urls + text_urls))

        # Core properties
        props = doc.core_properties

        return {
            "file_type":       "docx",
            "filename":        filename,
            "paragraph_count": len(doc.paragraphs),
            "extracted_text":  all_text.strip()[:5000],
            "extracted_urls":  all_urls,
            "url_count":       len(all_urls),
            "char_count":      len(all_text),
            "metadata": {
                "title":    props.title or "",
                "author":   props.author or "",
                "created":  str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
            },
            "ready_for_analysis": len(all_text.strip()) > 10,
        }

    except Exception as e:
        return {
            "file_type":   "docx",
            "filename":    filename,
            "error":       str(e),
            "extracted_text": "",
            "extracted_urls": [],
            "ready_for_analysis": False,
        }


# ── MASTER FORENSICS FUNCTION ──────────────────────────────────────────────────
async def run_forensics(file_bytes: bytes, filename: str, content_type: str) -> dict:
    """
    Main forensics function — routes to correct analyzer based on file type.
    """
    filename_lower = filename.lower()
    ext = Path(filename).suffix.lower()

    # Determine file type
    if ext in [".pdf"]:
        forensics_result = await analyze_pdf(file_bytes, filename)

    elif ext in [".docx", ".doc"]:
        forensics_result = await analyze_docx(file_bytes, filename)

    elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
        # Try QR first, then OCR
        qr_result = await analyze_qr(file_bytes, filename)
        if qr_result.get("qr_data"):
            forensics_result = qr_result
        else:
            forensics_result = await analyze_image(file_bytes, filename)

    else:
        return {
            "error": f"Unsupported file type: {ext}",
            "supported": [".png", ".jpg", ".jpeg", ".pdf", ".docx", ".webp"],
            "ready_for_analysis": False,
        }

    return {
        **forensics_result,
        "file_size_bytes": len(file_bytes),
        "file_size_kb":    round(len(file_bytes) / 1024, 2),
    }